#!/usr/bin/env python3
"""Generate per-product Instagram drafts and a daily posting reminder calendar."""

from __future__ import annotations

import json
import re
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PRODUCTS_PATH = ROOT / "frontend/public/demo-data/products.json"
ARTIFACTS = Path("/opt/cursor/artifacts")
PUBLIC_DOWNLOADS = ROOT / "frontend/public/downloads"
DOCX_PATH = ARTIFACTS / "Amit-Traders-Instagram-Product-Drafts.docx"
ICS_PATH = ARTIFACTS / "Amit-Traders-Daily-Post-Reminders.ics"
ICS_WEB_PATH = PUBLIC_DOWNLOADS / "instagram-post-reminders.ics"
MD_PATH = ARTIFACTS / "Amit-Traders-Instagram-Product-Drafts.md"
MD_WEB_PATH = PUBLIC_DOWNLOADS / "instagram-product-drafts.md"
SCHEDULE_PATH = ARTIFACTS / "Amit-Traders-Instagram-Posting-Schedule.csv"
SCHEDULE_WEB_PATH = PUBLIC_DOWNLOADS / "instagram-posting-schedule.csv"

SITE = "https://amittradersindia.com"
QUOTE_URL = f"{SITE}/quote"
API_BASE = "https://amit-traders-india-new.onrender.com"
INSTAGRAM = "@amittradersindia"
IST = ZoneInfo("Asia/Kolkata")

# Posting order: laptop bags first, then remaining categories.
CATEGORY_ORDER = [
    "Laptop Bags",
    "Gents Wallet",
    "Combo Set",
    "Duffle Bags",
    "Passport Holders",
    "Sling Bags",
]

CATEGORY_HASHTAGS = {
    "Laptop Bags": "#LaptopBag #ExecutiveBags #CorporateTravel",
    "Gents Wallet": "#LeatherWallet #MensWallet #WalletManufacturer",
    "Combo Set": "#CorporateCombo #GiftSet #CorporateGifting",
    "Duffle Bags": "#DuffleBag #TravelLeather #WeekendBag",
    "Passport Holders": "#PassportHolder #TravelAccessories",
    "Sling Bags": "#SlingBag #OfficeBag #CrossbodyBag",
}

BASE_HASHTAGS = (
    "#AmitTradersIndia #LeatherGoods #GenuineLeather #LeatherManufacturing "
    "#MadeInIndia #MumbaiBusiness #B2BLeather #CustomLeather #HandcraftedLeather "
    "#PrivateLabel #OEM #BulkOrders"
)

CAROUSEL_SLIDE_LABELS = [
    "Front view",
    "Side / profile view",
    "Detail / interior view",
    "Alternate angle",
    "Close-up craftsmanship",
]

# Weekdays: corporate audience scrolls during lunch (11:30 AM IST).
# Weekends: broader audience peaks in the evening (7:30 PM IST).
WEEKDAY_POST_HOUR = 11
WEEKDAY_POST_MINUTE = 30
WEEKEND_POST_HOUR = 19
WEEKEND_POST_MINUTE = 30
REMINDER_MINUTES_BEFORE = 15

START_DATE = date(2026, 8, 19)


def resolve_image_url(path: str | None) -> str:
    if not path:
        return ""
    if path.startswith("http"):
        return path
    if path.startswith("/uploads/"):
        return f"{API_BASE}{path}"
    return f"{SITE}{path}"


def clean_short_description(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    return "\n".join(line for line in lines if line)


def product_sort_key(product: dict) -> tuple:
    category = product["categoryName"]
    try:
        category_index = CATEGORY_ORDER.index(category)
    except ValueError:
        category_index = len(CATEGORY_ORDER)
    name = product.get("name", "")
    match = re.search(r"(\d+)$", name)
    number = int(match.group(1)) if match else 0
    return category_index, number, name


def load_products() -> list[dict]:
    data = json.loads(PRODUCTS_PATH.read_text(encoding="utf-8"))
    products = [p for p in data["content"] if p.get("active", True)]
    products.sort(key=product_sort_key)
    return products


def all_images(product: dict) -> list[str]:
    images = sorted(product.get("images") or [], key=lambda i: i.get("displayOrder", 0))
    return [resolve_image_url(img.get("imageUrl")) for img in images if img.get("imageUrl")]


def is_carousel(product: dict) -> bool:
    return len(all_images(product)) > 1


def post_format(product: dict) -> str:
    count = len(all_images(product))
    if count > 1:
        return f"Carousel ({count} slides)"
    return "Single image"


def posting_time_for_date(post_date: date) -> tuple[datetime, str]:
    is_weekday = post_date.weekday() < 5
    if is_weekday:
        post_at = datetime(
            post_date.year,
            post_date.month,
            post_date.day,
            WEEKDAY_POST_HOUR,
            WEEKDAY_POST_MINUTE,
            tzinfo=IST,
        )
        reason = "Weekday lunch scroll — best for corporate buyers and B2B decision-makers"
    else:
        post_at = datetime(
            post_date.year,
            post_date.month,
            post_date.day,
            WEEKEND_POST_HOUR,
            WEEKEND_POST_MINUTE,
            tzinfo=IST,
        )
        reason = "Weekend evening peak — higher general Instagram engagement in India"
    return post_at, reason


def format_post_time(post_at: datetime) -> str:
    return post_at.strftime("%-I:%M %p IST").replace("AM", "AM").replace("PM", "PM")


def carousel_slide_lines(product: dict) -> list[str]:
    images = all_images(product)
    lines = []
    for index, url in enumerate(images, start=1):
        label = CAROUSEL_SLIDE_LABELS[index - 1] if index <= len(CAROUSEL_SLIDE_LABELS) else f"View {index}"
        lines.append(f"Slide {index} ({label}): {url}")
    return lines


def build_caption(product: dict) -> str:
    short_desc = clean_short_description(product.get("shortDescription") or "")
    slug = product["slug"]
    category_slug = product["categorySlug"]
    product_url = f"{SITE}/products/{slug}"
    quote_with_product = f"{QUOTE_URL}?product={slug}"
    images = all_images(product)
    carousel = len(images) > 1

    lines = [
        f"{product['name']} | Amit Traders India",
        "",
    ]
    if carousel:
        lines.append(f"Swipe through all {len(images)} angles")
        lines.append("")
    lines.append(short_desc)
    lines.extend(
        [
            "",
            f"Category: {product['categoryName']}",
            f"Material: {product.get('material') or 'Leather'}",
        ]
    )
    if product.get("leatherType"):
        lines.append(f"Leather type: {product['leatherType']}")
    if product.get("colors"):
        lines.append(f"Colours: {product['colors']}")
    if product.get("minimumOrderQuantity"):
        lines.append(f"MOQ: {product['minimumOrderQuantity']} units")

    lines.extend(
        [
            "",
            "Custom branding available — embossing, foil stamping, and private label options.",
            "",
            "Request a bulk quote for this product:",
            quote_with_product,
            "",
            f"View product details: {product_url}",
            f"Browse {product['categoryName']}: {SITE}/products?category={category_slug}",
            "",
            f"{CATEGORY_HASHTAGS.get(product['categoryName'], '')} {BASE_HASHTAGS}".strip(),
        ]
    )
    return "\n".join(lines)


def build_alt_text(product: dict) -> str:
    short_desc = clean_short_description(product.get("shortDescription") or product["name"])
    return f"{product['name']} — {short_desc} by Amit Traders India"


def build_carousel_alt_texts(product: dict) -> list[str]:
    images = all_images(product)
    short_desc = clean_short_description(product.get("shortDescription") or product["name"])
    texts = []
    for index, _ in enumerate(images, start=1):
        label = CAROUSEL_SLIDE_LABELS[index - 1] if index <= len(CAROUSEL_SLIDE_LABELS) else f"view {index}"
        texts.append(f"{product['name']} — {label} — {short_desc} by Amit Traders India")
    return texts


def ics_escape(value: str) -> str:
    return (
        value.replace("\\", "\\\\")
        .replace(";", "\\;")
        .replace(",", "\\,")
        .replace("\n", "\\n")
    )


def build_ics(products: list[dict]) -> str:
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Amit Traders India//Instagram Post Reminders//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        "X-WR-CALNAME:Amit Traders Instagram Posts",
        "X-WR-TIMEZONE:Asia/Kolkata",
        "BEGIN:VTIMEZONE",
        "TZID:Asia/Kolkata",
        "BEGIN:STANDARD",
        "TZOFFSETFROM:+0530",
        "TZOFFSETTO:+0530",
        "TZNAME:IST",
        "DTSTART:19700101T000000",
        "END:STANDARD",
        "END:VTIMEZONE",
    ]

    for index, product in enumerate(products):
        post_date = START_DATE + timedelta(days=index)
        post_at, reason = posting_time_for_date(post_date)
        post_end = post_at + timedelta(minutes=20)
        uid = f"amit-ig-{product['slug']}-{post_date.isoformat()}@amittradersindia.com"
        name = product["name"]
        slug = product["slug"]
        short_desc = clean_short_description(product.get("shortDescription") or "")
        fmt = post_format(product)
        image_lines = "\\n".join(carousel_slide_lines(product)) if is_carousel(product) else all_images(product)[0]
        description = (
            f"Post {name} on {INSTAGRAM}\\n"
            f"Format: {fmt}\\n"
            f"Best time: {post_at.strftime('%I:%M %p IST')} — {reason}\\n\\n"
            f"Short description: {short_desc}\\n\\n"
            f"Images:\\n{image_lines}\\n\\n"
            f"Quote CTA: {QUOTE_URL}?product={slug}\\n\\n"
            f"Full drafts: {SITE}/downloads/instagram-product-drafts.md"
        )

        lines.extend(
            [
                "BEGIN:VEVENT",
                f"UID:{uid}",
                f"DTSTAMP:{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
                f"DTSTART;TZID=Asia/Kolkata:{post_at.strftime('%Y%m%dT%H%M%S')}",
                f"DTEND;TZID=Asia/Kolkata:{post_end.strftime('%Y%m%dT%H%M%S')}",
                f"SUMMARY:{ics_escape(f'Instagram {fmt}: {name}')}",
                f"DESCRIPTION:{ics_escape(description)}",
                "LOCATION:Instagram",
                "BEGIN:VALARM",
                "TRIGGER:-PT15M",
                "ACTION:DISPLAY",
                f"DESCRIPTION:{ics_escape(f'Reminder: post {name} at {post_at.strftime('%I:%M %p IST')}')}",
                "END:VALARM",
                "END:VEVENT",
            ]
        )

    lines.append("END:VCALENDAR")
    return "\r\n".join(lines) + "\r\n"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_engagement_guide(doc: Document) -> None:
    add_heading(doc, "Posting schedule & engagement guide", 1)
    doc.add_paragraph(
        "All 17 posts are scheduled from 19 August 2026. Import "
        "Amit-Traders-Daily-Post-Reminders.ics into your calendar for automatic reminders."
    )
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    guide = [
        ("Weekdays (Mon–Fri)", "Post at 11:30 AM IST — reminder at 11:15 AM IST"),
        ("Weekends (Sat–Sun)", "Post at 7:30 PM IST — reminder at 7:15 PM IST"),
        ("Why weekdays at lunch", "Corporate buyers and procurement teams browse during mid-morning breaks"),
        ("Why weekends in evening", "Higher leisure scrolling and discovery engagement across India"),
        ("Carousel posts", "Laptop Bag 005–009 use all product images as swipeable carousel slides"),
    ]
    for i, (label, value) in enumerate(guide):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    doc.add_paragraph()
    doc.add_paragraph(
        "Tip: Reply to DMs within 1 hour of posting and add the quote link as the first comment "
        f"({QUOTE_URL}) for extra clicks."
    )
    doc.add_page_break()


def build_docx(products: list[dict]) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Amit Traders India\nInstagram Product Post Drafts")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3D, 0x2B, 0x1F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    carousel_count = sum(1 for p in products if is_carousel(p))
    sub = subtitle.add_run(
        f"{len(products)} daily posts ({carousel_count} carousels) • {INSTAGRAM} • "
        f"Starting {START_DATE.strftime('%d %B %Y')}"
    )
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_engagement_guide(doc)

    current_category = None
    day = 0

    for product in products:
        day += 1
        post_date = START_DATE + timedelta(days=day - 1)
        post_at, reason = posting_time_for_date(post_date)
        category = product["categoryName"]
        images = all_images(product)
        carousel = is_carousel(product)

        if category != current_category:
            if current_category is not None:
                doc.add_page_break()
            current_category = category
            add_heading(doc, category, 1)

        add_heading(doc, f"Day {day} — {post_date.strftime('%a, %d %b %Y')} — {product['name']}", 2)

        table = doc.add_table(rows=8, cols=2)
        table.style = "Table Grid"
        rows = [
            ("Scheduled date", post_date.strftime("%A, %d %B %Y")),
            ("Post at", f"{post_at.strftime('%I:%M %p IST')} ({reason})"),
            ("Format", post_format(product)),
            ("Product", product["name"]),
            ("Short description", clean_short_description(product.get("shortDescription") or "")),
            ("Product page", f"{SITE}/products/{product['slug']}"),
            ("Quote CTA", f"{QUOTE_URL}?product={product['slug']}"),
            ("Instagram steps", "Create → Post → Carousel (if applicable) → upload slides in order → paste caption"),
        ]
        for i, (label, value) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_paragraph()
        if carousel:
            add_heading(doc, f"Carousel slides ({len(images)} images — upload in this order)", 3)
            slide_table = doc.add_table(rows=len(images) + 1, cols=3)
            slide_table.style = "Table Grid"
            slide_table.rows[0].cells[0].text = "Slide"
            slide_table.rows[0].cells[1].text = "Label"
            slide_table.rows[0].cells[2].text = "Image URL"
            for index, url in enumerate(images, start=1):
                label = CAROUSEL_SLIDE_LABELS[index - 1] if index <= len(CAROUSEL_SLIDE_LABELS) else f"View {index}"
                slide_table.rows[index].cells[0].text = str(index)
                slide_table.rows[index].cells[1].text = label
                slide_table.rows[index].cells[2].text = url
            doc.add_paragraph()
        else:
            add_heading(doc, "Image", 3)
            doc.add_paragraph(images[0] if images else "No image on file")

        add_heading(doc, "Caption (copy & paste)", 3)
        caption = doc.add_paragraph()
        caption.add_run(build_caption(product)).italic = True

        add_heading(doc, "Alt text", 3)
        if carousel:
            for index, alt in enumerate(build_carousel_alt_texts(product), start=1):
                doc.add_paragraph(f"Slide {index}: {alt}")
        else:
            doc.add_paragraph(build_alt_text(product))

        doc.add_paragraph()

    return doc


def build_markdown(products: list[dict]) -> str:
    carousel_count = sum(1 for p in products if is_carousel(p))
    lines = [
        "# Amit Traders India — Instagram Product Post Drafts",
        "",
        f"- **Account:** {INSTAGRAM}",
        f"- **Products:** {len(products)} ({carousel_count} carousel posts)",
        f"- **Schedule starts:** {START_DATE.strftime('%d %B %Y')} (daily)",
        f"- **Reminder file:** `Amit-Traders-Daily-Post-Reminders.ics`",
        "",
        "## Best times to post (India / IST)",
        "",
        "| Day type | Post at | Reminder at | Why |",
        "|----------|---------|-------------|-----|",
        "| Weekdays (Mon–Fri) | **11:30 AM IST** | 11:15 AM IST | Corporate buyers scroll during lunch breaks |",
        "| Weekends (Sat–Sun) | **7:30 PM IST** | 7:15 PM IST | Higher evening engagement across India |",
        "",
        "**Carousel posts:** Laptop Bag 005, 006, 007, 008, 009 — upload all slides in order.",
        "",
        "---",
        "",
    ]

    current_category = None
    day = 0

    for product in products:
        day += 1
        post_date = START_DATE + timedelta(days=day - 1)
        post_at, reason = posting_time_for_date(post_date)
        category = product["categoryName"]
        images = all_images(product)
        carousel = is_carousel(product)

        if category != current_category:
            current_category = category
            lines.extend([f"## {category}", ""])

        lines.extend(
            [
                f"### Day {day} — {post_date.strftime('%a, %d %b %Y')} — {product['name']}",
                "",
                f"**Post at:** {post_at.strftime('%I:%M %p IST')} — {reason}",
                f"**Format:** {post_format(product)}",
                f"**Short description:** {clean_short_description(product.get('shortDescription') or '')}",
                "",
            ]
        )

        if carousel:
            lines.append("**Carousel slides (upload in order):**")
            lines.append("")
            for slide_line in carousel_slide_lines(product):
                lines.append(f"- {slide_line}")
            lines.append("")
        else:
            lines.extend([f"**Image:** {images[0] if images else 'N/A'}", ""])

        lines.extend(
            [
                f"**Quote CTA:** {QUOTE_URL}?product={product['slug']}",
                "",
                "```",
                build_caption(product),
                "```",
                "",
            ]
        )

        if carousel:
            lines.append("**Alt text (per slide):**")
            lines.append("")
            for index, alt in enumerate(build_carousel_alt_texts(product), start=1):
                lines.append(f"- Slide {index}: {alt}")
            lines.append("")
        else:
            lines.append(f"**Alt text:** {build_alt_text(product)}")
            lines.append("")

        lines.extend(["---", ""])

    return "\n".join(lines)


def build_schedule_csv(products: list[dict]) -> str:
    rows = [
        "Day,Date,Post Time (IST),Format,Product,Category,Short Description,Quote CTA,Image Count"
    ]
    for index, product in enumerate(products, start=1):
        post_date = START_DATE + timedelta(days=index - 1)
        post_at, _ = posting_time_for_date(post_date)
        short_desc = clean_short_description(product.get("shortDescription") or "").replace('"', '""')
        rows.append(
            f'{index},{post_date.isoformat()},{post_at.strftime("%I:%M %p IST")},'
            f'{post_format(product)},"{product["name"]}","{product["categoryName"]}",'
            f'"{short_desc}",{QUOTE_URL}?product={product["slug"]},{len(all_images(product))}'
        )
    return "\n".join(rows) + "\n"


def main() -> None:
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    PUBLIC_DOWNLOADS.mkdir(parents=True, exist_ok=True)
    products = load_products()

    markdown = build_markdown(products)
    ics = build_ics(products)
    schedule = build_schedule_csv(products)

    build_docx(products).save(DOCX_PATH)
    ICS_PATH.write_text(ics, encoding="utf-8")
    ICS_WEB_PATH.write_text(ics, encoding="utf-8")
    MD_PATH.write_text(markdown, encoding="utf-8")
    MD_WEB_PATH.write_text(markdown, encoding="utf-8")
    SCHEDULE_PATH.write_text(schedule, encoding="utf-8")
    SCHEDULE_WEB_PATH.write_text(schedule, encoding="utf-8")

    carousel_products = [p["name"] for p in products if is_carousel(p)]
    print(f"Products: {len(products)}")
    print(f"Carousels: {len(carousel_products)} ({', '.join(carousel_products)})")
    print(f"Saved: {DOCX_PATH}")
    print(f"Saved: {ICS_PATH}")
    print(f"Saved: {ICS_WEB_PATH}")
    print(f"Saved: {MD_PATH}")
    print(f"Saved: {MD_WEB_PATH}")
    print(f"Saved: {SCHEDULE_PATH}")
    print(f"Saved: {SCHEDULE_WEB_PATH}")


if __name__ == "__main__":
    main()
